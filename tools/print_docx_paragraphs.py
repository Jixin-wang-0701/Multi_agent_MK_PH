from docx import Document
import sys


sys.stdout.reconfigure(encoding="utf-8")

document = Document(sys.argv[1])
start = int(sys.argv[2]) if len(sys.argv) > 2 else 0
stop = int(sys.argv[3]) if len(sys.argv) > 3 else len(document.paragraphs)
for index, paragraph in enumerate(document.paragraphs[start:stop], start=start):
    text = paragraph.text.strip()
    if text:
        print(index, text[:1000])
