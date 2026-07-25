from pathlib import Path
import shutil

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript相关" / "Revised MS_ZJW_NCB_Fig6.docx"
OUTPUT = ROOT / "manuscript相关" / "Revised MS_ZJW_NCB_Fig6_revised.docx"


REPLACEMENTS = {
    32: (
        "Determining how hypoxic lung megakaryocytes (MKs) reshape vascular disease requires the integration of cell-resolved transcriptomics, metabolomics, prior in vivo phenotypes and knowledge resources; conventional serial analyses do not readily rank competing causal models across these evidence layers. "
        "We therefore developed a multi-agent AI system designed to generate, critique and prioritize explicitly evidence-traceable, experimentally falsifiable hypotheses rather than a black-box prediction (Fig. 6A,B)."
    ),
    33: (
        "The system separates governance, evidence assembly and adversarial evaluation into specialized, coordinated modules (Fig. 6A,B). "
        "The PI Agent formulates the biological question, prespecifies evidence standards and defines the research brief. Bioinformatics, metabolic and tool-use modules assemble a shared, auditable evidence context from our scRNA-seq and metabolomics data, prior in vivo findings, pathway resources, literature and public-dataset analyses. Independent Generation Agents then propose non-redundant mechanism-level directions. The Proximity Check Agent merges overlap, Reflection Agents challenge mechanistic plausibility, evidence quality, feasibility, falsifiability and overinterpretation, and the Ranking Agent compares candidates using convergent support. Finally, Meta-review and Evolution Agents synthesize recurrent weaknesses, refine the strongest directions and return them to the PI Agent for an explicit advance, revise, merge or reject decision. This iterative design preserves the evidence and critique associated with each decision, so a hypothesis advances only when it is sufficiently supported, coherent, novel and testable; detailed prompts and algorithmic implementation are described in Methods."
    ),
    34: (
        "Across independent generation branches, the system converged on an AMD1-centered metabolic direction as the highest-priority hypothesis (Fig. 6C). Three orthogonal data streams supported this prioritization. First, methionine abundance was increased in PH-MK samples relative to control MK samples (log2FC = 3.26). Second, Amd1 was detected in 31.44% of MKs versus 14.87% of all other cells and was enriched in MKs (log2 enrichment = 1.353). Third, Amd1 expression was higher in PH MKs than in control MKs (log2FC = 1.77; Wilcoxon P = 6.55e-06). The system therefore linked the methionine signal to Amd1 as a pathway-neighbor anchor within methionine/S-adenosylmethionine/polyamine metabolism; this association supports prioritization but does not by itself establish a direct enzymatic or causal connection. Competing candidates were down-ranked because their support was less convergent: Amd2 showed substantially lower MK expression, whereas Dnmt3b and Cyp26b1 lacked a robust PH-associated MK shift. The resulting direction proposes that hypoxic lung MKs acquire an AMD1-linked methionine/S-adenosylmethionine/polyamine state that can influence pro-remodeling immune, vascular-wall or extracellular-vesicle/stromal programs. These downstream routes were retained as alternatives to be discriminated experimentally, rather than presented as a settled bridge."
    ),
    35: (
        "The prioritized output specified AMD1 inhibition or MK-selective Amd1 perturbation as the key falsification experiment, coupled to targeted LC-MS quantification of methionine/S-adenosylmethionine/polyamines, immune and vascular readouts, and vascular-remodeling phenotypes. These predictions are tested in the subsequent validation section, closing the loop from AI-guided hypothesis prioritization to experimental evaluation (Fig. 6C)."
    ),
    67: "Fig. 6 | Multi-agent AI system prioritizes an AMD1-centered metabolic direction linking hypoxic lung MKs to vascular remodeling.",
    68: (
        "(A) Architecture of the multi-agent co-scientist framework. The system receives the biological question, prior in vivo MK evidence and multi-omics inputs, then separates PI-guided planning, bioinformatic and metabolic evidence assembly, hypothesis generation, critical review, ranking, synthesis and evolution into specialized modules. The output is a ranked, evidence-traceable set of experimentally testable directions."
    ),
    69: (
        "(B) One-cycle workflow. A PI-defined brief and shared evidence context feed parallel bioinformatic, metabolic and hypothesis-generation branches. Proximity, reflection and ranking modules then independently de-duplicate, critique and prioritize candidate directions before meta-review/evolution and PI feedback determine the next cycle."
    ),
    70: (
        "(C) Evidence convergence for the prioritized AMD1 direction. The system integrates increased methionine in PH-MK samples (log2FC = 3.26), MK-enriched Amd1 expression (31.44% of MKs versus 14.87% of other cells; log2 enrichment = 1.353), and Amd1 upregulation in PH MKs (log2FC = 1.77; Wilcoxon P = 6.55e-06). These data prioritize an AMD1-linked methionine/S-adenosylmethionine/polyamine state and nominate AMD1 perturbation with metabolic, immune, vascular and morphometric readouts as a falsification strategy; the downstream immune, vascular-wall and extracellular-vesicle/stromal routes remain candidate mechanisms."
    ),
}


def replace_paragraph(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


shutil.copy2(SOURCE, OUTPUT)
document = Document(OUTPUT)
for index, replacement in REPLACEMENTS.items():
    replace_paragraph(document.paragraphs[index], replacement)
document.save(OUTPUT)
print(OUTPUT)
