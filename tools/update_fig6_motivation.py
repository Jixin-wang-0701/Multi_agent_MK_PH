from pathlib import Path

from docx import Document


OLD = (
    "Determining how hypoxic lung megakaryocytes (MKs) reshape vascular disease requires the integration of cell-resolved transcriptomics, metabolomics, prior in vivo phenotypes and knowledge resources; conventional serial analyses do not readily rank competing causal models across these evidence layers. "
    "We therefore developed a multi-agent AI system designed to generate, critique and prioritize explicitly evidence-traceable, experimentally falsifiable hypotheses rather than a black-box prediction (Fig. 6A,B)."
)

NEW = (
    "Uncovering how hypoxic lung megakaryocytes (MKs) reshape vascular disease is a genuinely multi-layered causal-inference problem: cell-resolved transcriptional states, MK-enriched metabolite shifts, prior in vivo phenotypes and external knowledge each nominate partially overlapping, often competing mechanisms, whereas conventional serial analyses usually interrogate these evidence layers separately and provide no transparent way to challenge and rank the resulting causal models. "
    "We therefore developed a multi-agent AI system in which specialized agents independently generate mechanistic hypotheses, retrieve and assemble supporting evidence, interrogate alternative explanations and critically review one another's outputs before prioritizing only evidence-traceable, experimentally falsifiable models; this collaborative, auditable design reduces dependence on any single analytical perspective and moves beyond black-box prediction to a rigorous route from heterogeneous data to actionable causal tests (Fig. 6A,B)."
)


def main() -> None:
    candidates = list(Path.cwd().rglob("Revised MS_ZJW_NCB_Fig6_revised.docx"))
    if len(candidates) != 1:
        raise RuntimeError(f"Expected one revised manuscript, found: {candidates}")

    path = candidates[0]
    document = Document(path)
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == OLD]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one target paragraph, found {len(matches)}")

    paragraph = matches[0]
    paragraph.clear()
    paragraph.add_run(NEW)
    output = path.with_name("Revised MS_ZJW_NCB_Fig6_revised_motivation.docx")
    document.save(output)
    print(output)


if __name__ == "__main__":
    main()
